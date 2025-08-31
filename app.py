import os, io, json, re
from datetime import datetime
import streamlit as st

st.set_page_config(page_title="InAA — No-Login Demo", page_icon="🧭", layout="wide")

# --- Light, safe startup: no heavy imports yet ---
API_KEY = os.environ.get("OPENAI_API_KEY", "").strip()

st.title("InAA — Inclusive Apprenticeship Assistant (No-Login Demo)")
st.caption("Upload, lint, chat, and export. No ChatGPT account required.")

# Status panel (helps diagnose blanks)
with st.expander("Status & Diagnostics"):
    st.write("Python:", os.sys.version)
    st.write("OpenAI API:", "Enabled ✅" if API_KEY else "Off")
    st.write("Repo contents:", os.listdir("."))
    st.write("Assets present:", os.path.isdir("assets"),
             os.path.exists("assets/linter_rules_v01.json"))

# --- Lazy loaders to avoid hard crashes on import ---
def try_imports(for_pdf=False, for_docx=False, for_excel=False, for_openai=False, for_pandas=False):
    mods = {}
    try:
        if for_pdf:
            from pypdf import PdfReader
            mods["PdfReader"] = PdfReader
    except Exception as e:
        st.warning(f"PDF support unavailable: {e}")
    try:
        if for_docx:
            from docx import Document as DocxDocument
            mods["DocxDocument"] = DocxDocument
    except Exception as e:
        st.warning(f"DOCX support unavailable: {e}")
    try:
        if for_excel or for_pandas:
            import pandas as pd
            mods["pd"] = pd
    except Exception as e:
        st.warning(f"Pandas not available: {e}")
    try:
        if for_excel:
            import openpyxl  # noqa: F401
            mods["openpyxl"] = True
    except Exception as e:
        st.warning(f"Excel export unavailable: {e}")
    try:
        if for_openai and API_KEY:
            from openai import OpenAI
            mods["OpenAI"] = OpenAI
    except Exception as e:
        st.warning(f"OpenAI client not available: {e}")
    return mods

# --- Knowledge (optional) ---
def load_knowledge():
    texts = []
    if os.path.isdir("assets/knowledge"):
        for name in os.listdir("assets/knowledge"):
            if name.endswith(".txt"):
                try:
                    texts.append(open(os.path.join("assets/knowledge", name), "r", encoding="utf-8").read())
                except:
                    pass
    return texts

KNOWLEDGE = load_knowledge()

# --- Sidebar: Guided Starters ---
with st.sidebar:
    st.header("Guided Starters")
    starter = st.radio("Choose a task:", [
        "Draft my WPS", "Audit this WPS/RTI", "Accommodation SOP", "Partner & Funding map"
    ])

    if starter == "Draft my WPS":
        with st.form("form_wps"):
            occ = st.text_input("Occupation", placeholder="IT Support Specialist")
            region = st.text_input("Region/County", placeholder="San Mateo County")
            structure = st.selectbox("Structure", ["Hybrid","Competency","Time-based"], index=0)
            union = st.text_input("Union/Labor-Management (optional)", placeholder="None")
            rti = st.text_input("RTI Provider (optional)", placeholder="Local Community College")
            submitted = st.form_submit_button("Create Draft Prompt")
        if submitted:
            seed = (f"Draft a {structure.lower()} Work Process Schedule for {occ} in {region}. "
                    f"Use task-not-method verbs, include accessibility notes, and suggest competencies. "
                    f"Union: {union or 'N/A'}. RTI: {rti or 'N/A'}.")
            st.session_state["seed"] = seed

    elif starter == "Audit this WPS/RTI":
        st.info("Upload or paste your draft below; then click **Run Accessibility Linter**.")
        st.session_state["seed"] = "Audit the provided draft for accessibility issues, show flags, and produce a clean copy."

    elif starter == "Accommodation SOP":
        with st.form("form_sop"):
            org = st.text_input("Employer/Program (optional)", placeholder="Small HVAC contractor")
            sla_days = st.number_input("SLA (days to initial meeting)", min_value=1, max_value=30, value=10)
            submitted = st.form_submit_button("Create SOP Prompt")
        if submitted:
            st.session_state["seed"] = (
                f"Create an Accommodation SOP for {org or 'a small employer'} with community-college RTI, "
                f"include a {sla_days}-day SLA and privacy notes. Provide scripts and a request workflow."
            )

    else:  # Partner & Funding map
        with st.form("form_pf"):
            county = st.text_input("County/Region", placeholder="Los Angeles County")
            pop = st.text_input("Target population (optional)", placeholder="Veterans, neurodivergent learners")
            submitted = st.form_submit_button("Create Partner/Funding Prompt")
        if submitted:
            st.session_state["seed"] = (
                f"Draft a partner map (DOR/VR, AJCs, CILs, unions, CBOs, education) and a braided funding sketch "
                f"for {county}. Target population: {pop or 'General apprenticeship candidates'}."
            )

    st.divider()
    st.write("OpenAI API:", "Enabled ✅" if API_KEY else "Off (linter still works)")

# --- Upload / Paste ---
left, right = st.columns(2)
with left:
    up = st.file_uploader("Upload WPS/RTI (.docx, .pdf, .txt, .md)", type=["docx","pdf","txt","md"])
with right:
    pasted = st.text_area("Or paste your text here", height=220, placeholder="Paste any WPS/RTI text here…")

text = ""
if up is not None:
    name = up.name.lower()
    data = up.read()
    if name.endswith((".txt",".md")):
        text = data.decode("utf-8", "ignore")
    elif name.endswith(".docx"):
        mods = try_imports(for_docx=True)
        DocxDocument = mods.get("DocxDocument")
        if DocxDocument:
            doc = DocxDocument(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join([c.text for c in row.cells]))
            text = "\n".join(parts)
        else:
            st.error("DOCX support not available. Please paste text.")
    elif name.endswith(".pdf"):
        mods = try_imports(for_pdf=True)
        PdfReader = mods.get("PdfReader")
        if PdfReader:
            reader = PdfReader(io.BytesIO(data))
            pages = []
            for page in reader.pages:
                try:
                    pages.append(page.extract_text() or "")
                except:
                    pages.append("")
            text = "\n".join(pages)
        else:
            st.error("PDF support not available. Please paste text.")
else:
    if pasted.strip():
        text = pasted

# --- Seed into chat if present ---
st.divider()
st.subheader("Ask Questions / Run Starters")
if "chat" not in st.session_state:
    st.session_state["chat"] = []
if "seed" in st.session_state:
    st.chat_message("assistant").write(st.session_state["seed"])
    if API_KEY:
        mods = try_imports(for_openai=True)
        OpenAI = mods.get("OpenAI")
        if OpenAI:
            client = OpenAI(api_key=API_KEY)
            context = "\n\n".join([k[:1800] for k in KNOWLEDGE[:3]])
            try:
                msg = st.session_state["seed"] + ("\n\nContext:\n" + context if context else "")
                resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role":"system","content":"You are the Accessibility WPS Assistant. Use inclusive, plain language."},
                              {"role":"user","content":msg}],
                    temperature=0.3, max_tokens=800
                )
                st.chat_message("assistant").write(resp.choices[0].message.content)
            except Exception as e:
                st.warning(f"Chat unavailable: {e}")
    del st.session_state["seed"]

user_msg = st.chat_input("Ask about WPS/RTI, accessibility, partners, or funding…")
if user_msg:
    st.chat_message("user").write(user_msg)
    if API_KEY:
        mods = try_imports(for_openai=True)
        OpenAI = mods.get("OpenAI")
        if OpenAI:
            client = OpenAI(api_key=API_KEY)
            context = "\n\n".join([k[:1800] for k in KNOWLEDGE[:3]])
            try:
                resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role":"system","content":"You are the Accessibility WPS Assistant. Use inclusive, plain language."},
                              {"role":"user","content":user_msg + ("\n\nContext:\n"+context if context else "")}],
                    temperature=0.3, max_tokens=800
                )
                st.chat_message("assistant").write(resp.choices[0].message.content)
            except Exception as e:
                st.warning(f"Chat unavailable: {e}")
    else:
        st.chat_message("assistant").write("AI chat is off (no API key). I can still lint text and export files.")

# --- Linter (uses only stdlib regex) ---
st.divider()
st.subheader("Accessibility Linter")
flags, clean = [], ""
if st.button("Run Accessibility Linter"):
    # Minimal rules inline to avoid JSON issues
    RULES = [
        {"id":"R-A1","pattern":r"\bclimb(ing)?\b","flags":"i","severity":"warn",
         "message":"Replace physical-method verb with task-focused wording.",
         "suggestion":"Use 'ascend a ladder' or 'access elevated work areas'."},
        {"id":"R-A2","pattern":r"lift(?:ing)?\s*(\d+)\s*(lb|lbs|pounds|kg)?","flags":"i","severity":"warn",
         "message":"Hard physical requirement may exclude qualified candidates if not essential.",
         "suggestion":"State the task and allow assistive devices/team lifts."},
        {"id":"R-B3","pattern":r"(excellent communication|team player|self[- ]starter|strong work ethic|detail[- ]oriented)","flags":"i","severity":"info",
         "message":"Vague soft-skill language.","suggestion":"Define observable behaviors."},
        {"id":"R-D1","pattern":r"with or without reasonable accommodation","flags":"i","severity":"info",
         "message":"ADA boilerplate belongs in HR, not WPS.","suggestion":"Remove from WPS; keep in policy docs."},
    ]
    if not text.strip():
        st.warning("Upload or paste some text first.")
    else:
        for r in RULES:
            flg = re.I if "i" in r["flags"] else 0
            for m in re.finditer(r["pattern"], text, flg):
                flags.append({
                    "rule_id": r["id"], "severity": r["severity"],
                    "match": m.group(0), "message": r["message"], "suggestion": r["suggestion"]
                })
        clean = re.sub(r"\bclimb(ing)?\b","ascend", text, flags=re.I)
        clean = re.sub(r"lift(?:ing)?\s*(\d+)\s*(lb|lbs|pounds|kg)?",
                       r"move materials up to \1 \2 using safe methods", clean, flags=re.I)
        clean = re.sub(r"(excellent communication|team player|self[- ]starter|strong work ethic|detail[- ]oriented)",
                       r"communicates clearly with mentors and closes the loop on tasks", clean, flags=re.I)

        st.success(f"Found {len(flags)} potential issues.")
        if flags:
            try:
                import pandas as pd
                st.dataframe(pd.DataFrame(flags))
            except Exception:
                st.write(flags)
        st.subheader("Clean Copy (rule-based)")
        st.write(clean[:4000])

        if st.checkbox("Polish with AI (if API enabled)") and API_KEY:
            mods = try_imports(for_openai=True)
            OpenAI = mods.get("OpenAI")
            if OpenAI:
                client = OpenAI(api_key=API_KEY)
                try:
                    resp = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role":"system","content":"Rewrite into inclusive, task-focused language; plain style."},
                                  {"role":"user","content":clean}],
                        temperature=0.3, max_tokens=1200
                    )
                    clean = resp.choices[0].message.content
                    st.subheader("Clean Copy (polished)")
                    st.write(clean[:4000])
                except Exception as e:
                    st.warning(f"Polish unavailable: {e}")

# --- Exports (lazy import python-docx / pandas only when needed) ---
st.divider()
st.subheader("Exports")
c1, c2, c3 = st.columns(3)

with c1:
    if st.button("Export WPS (.docx)"):
        body = clean if clean else text
        if not body.strip():
            st.warning("Run the linter first (or paste/upload some text).")
        else:
            mods = try_imports(for_docx=True)
            DocxDocument = mods.get("DocxDocument")
            if not DocxDocument:
                st.error("DOCX export unavailable (python-docx not installed).")
            else:
                # Create a simple doc from scratch if template missing
                try:
                    from docx import Document as DocxDocument  # ensure import
                    doc = DocxDocument()
                    doc.add_heading("Linted/Rewritten WPS", 0)
                    for line in body.split("\n"):
                        doc.add_paragraph(line)
                    bio = io.BytesIO()
                    doc.save(bio); bio.seek(0)
                    st.download_button("Download WPS", data=bio,
                        file_name=f"WPS_Clean_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
                except Exception as e:
                    st.error(f"Export failed: {e}")

with c2:
    if st.button("Export RTI (.docx)"):
        notes = "RTI Outline (Derived)\n- Modules with UDL hooks\n- Accessible materials\n- Performance-based assessments\n"
        if clean.strip(): notes += "\nNotes:\n" + clean[:1500]
        mods = try_imports(for_docx=True)
        DocxDocument = mods.get("DocxDocument")
        if DocxDocument:
            doc = DocxDocument(); doc.add_heading("RTI Outline", 0)
            for line in notes.split("\n"): doc.add_paragraph(line)
            bio = io.BytesIO(); doc.save(bio); bio.seek(0)
            st.download_button("Download RTI", data=bio,
                file_name=f"RTI_Outline_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
        else:
            st.error("DOCX export unavailable (python-docx not installed).")

with c3:
    if st.button("Export Checklist (.xlsx)"):
        mods = try_imports(for_excel=True, for_pandas=True)
        pd = mods.get("pd")
        if not pd:
            st.error("Excel export unavailable (pandas/openpyxl not installed).")
        else:
            df = pd.DataFrame(flags or [{"rule_id":"","severity":"","match":"","message":"","suggestion":""}])
            bio = io.BytesIO()
            with pd.ExcelWriter(bio, engine="openpyxl") as w:
                df.to_excel(w, index=False, sheet_name="Linter Flags")
            bio.seek(0)
            st.download_button("Download Checklist", data=bio,
                file_name=f"Accessibility_Checklist_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx")

st.caption("Tip: Prefer DOCX or paste text if a PDF's text layer is messy.")
